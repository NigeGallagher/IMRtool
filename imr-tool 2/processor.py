"""
IMR contributor submission processor.

This deliberately does NOT hand-build IDML XML. The earlier version of this
tool spent a long debugging session fighting InDesign over hand-rolled IDML
(ID clashes, missing required Spread/Page attributes, font inheritance via
Properties/AppliedFont, etc) and still came out the other side with fonts
not resolving correctly.

Instead, this produces a clean .docx with the IMR paragraph styles already
named and formatted (Article Title, Byline, Standfirst, Body Text First,
Body Text, Subhead, Pull Quote, Caption, Endnote Text, Footer). When you
File > Place that .docx into an InDesign document that already has
paragraph styles with those same names, InDesign maps the incoming text to
your existing styles automatically (tick "Preserve Styling" / use the style
mapping dialog on place). No raw IDML to fight with, and you can always
open the .docx yourself to sanity-check the content before it ever touches
InDesign.
"""

import os
import zipfile
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, Mm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
NSMAP = {'w': W_NS}

# ─── IMR typographic system ───
# (font, size, leading, bold, italic, all_caps, justify, space_after_pt)
STYLES = {
    'Article Title':   ('Barlow Condensed', 36, 34, True,  False, True,  False, 0),
    'Byline':          ('Barlow Condensed', 12, 14, True,  False, True,  False, 6),
    'Standfirst':      ('Source Serif 4',   13, 18, True,  False, False, False, 8),
    'Body Text First': ('Source Serif 4',   10, 15, False, False, False, True,  8),
    'Body Text':       ('Source Serif 4',   10, 15, False, False, False, True,  8),
    'Subhead':         ('Barlow Condensed', 13, 16, True,  False, True,  False, 6),
    'Pull Quote':      ('Barlow Condensed', 18, 20, True,  False, True,  False, 8),
    'Caption':         ('Source Serif 4',    9, 12, False, True,  False, False, 6),
    'Endnote Text':    ('Source Serif 4',    9, 12, False, False, False, False, 4),
    'Footer':          ('Barlow Condensed',  9, 11, True,  False, True,  False, 0),
}

MARKER_STYLE = {
    '[PULLQUOTE]': 'Pull Quote',
    '[SUBHEAD]':   'Subhead',
    '[CAPTION]':   'Caption',
    '[ENDNOTE]':   'Endnote Text',
}

# Page size - the journal's actual trim size, not a standard Word
# default. Applied to both the single-column intro section and the
# two-column body section.
PAGE_WIDTH_MM = 200
PAGE_HEIGHT_MM = 250

# Page margins, applied to both the single-column intro section and the
# two-column body section. Eased back partway from a first pass that ran
# too tight - still noticeably slimmer than Word's US default (1"/1.25")
# but with a bit more breathing room than the initial half-default cut.
MARGIN_TOP_IN = 0.7
MARGIN_BOTTOM_IN = 0.7
MARGIN_LEFT_IN = 0.85
MARGIN_RIGHT_IN = 0.85

# Rough word budget for the single-column intro page before switching to
# the two-column section. This needs to leave room for IMAGE_GAP_PT below
# - that fixed-height gap eats roughly 12 lines (180pt / 15pt leading) at
# around 14 words a line, so the word target is set well under what would
# fit without the gap. This is a heuristic, not a guarantee of exactly
# filling page 1 - actual fit depends on how Word/InDesign renders the
# fonts, and on how long the title and standfirst happen to be. Tune this
# number (and IMAGE_GAP_PT below) together if the intro page is running
# short or long in practice.
INTRO_WORD_TARGET = 70

# Styles that flow as ordinary running body copy and should stop the
# single-column intro section once the word target is hit. Subheads,
# pull quotes, captions etc. always start the two-column section even if
# the word target hasn't been reached yet, since that's the natural
# break point in print layout.
INTRO_ELIGIBLE_STYLES = {'Body Text First', 'Body Text'}

# Heuristic for catching section headings that the contributor didn't
# explicitly tag with [SUBHEAD] or a Word Heading style. In practice,
# real flowing prose from a Word document almost always ends in actual
# punctuation - contributors write complete sentences - so "short, with
# no terminal punctuation" is already doing most of the real work of
# telling a heading apart from a sentence. Earlier versions of this also
# required every word to be capitalised (Title Case, e.g. "Silicon
# Economy"), which missed plenty of real headings that use sentence case
# instead (e.g. "Critique of the family") - so capitalisation here is
# just a basic sanity check, not the main signal.
MAX_SUBHEAD_WORDS = 8

# Fixed blank space left on page 1, between the standfirst and the start
# of the body text - reserved for a photo or illustration to be dropped
# in later in InDesign. This is a flat point value rather than anything
# based on content length, so every article gets the same gap regardless
# of how long the standfirst or lead paragraph is.
IMAGE_GAP_PT = 180  # roughly 2.5 inches


def _set_all_caps(style, value):
    """python-docx doesn't expose all_caps on every version cleanly via the
    high level API in all paragraph-style contexts, so set it directly on
    the underlying rPr to be safe."""
    rpr = style.element.get_or_add_rPr()
    caps = rpr.find(qn('w:caps'))
    if value:
        if caps is None:
            caps = rpr.makeelement(qn('w:caps'), {})
            rpr.append(caps)
    else:
        if caps is not None:
            rpr.remove(caps)


def _set_columns(section, num, space_twips=360):
    """Set the column count on a section. python-docx always creates a
    default <w:cols> element in the correct schema position (after pgMar,
    before docGrid), so we just set attributes on the existing one rather
    than inserting a new element and risking putting it in the wrong spot."""
    sectPr = section._sectPr
    cols = sectPr.find(qn('w:cols'))
    cols.set(qn('w:num'), str(num))
    cols.set(qn('w:space'), str(space_twips))


def _set_page_size(section):
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)


def _set_margins(section):
    section.top_margin = Inches(MARGIN_TOP_IN)
    section.bottom_margin = Inches(MARGIN_BOTTOM_IN)
    section.left_margin = Inches(MARGIN_LEFT_IN)
    section.right_margin = Inches(MARGIN_RIGHT_IN)


def _add_fixed_gap(doc, points):
    """Insert an empty paragraph with an exact, fixed line height - a
    deterministic blank space that's the same regardless of font,
    content, or anything else, used to reserve room for an image."""
    gap = doc.add_paragraph()
    pf = gap.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(points)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    return gap


# Note markers are sized relative to the body text they sit in, scaled
# down to roughly footnote-reference size - standard practice is to make
# the superscript number noticeably smaller than the running text rather
# than just raising it at full size.
NOTE_MARKER_SCALE = 0.7


def _add_paragraph_with_notes(doc, segments, style_name):
    """Add a paragraph built from a list of ('text', string) and
    ('note', [numbers]) segments, in order - so a footnote/endnote
    reference renders as a small, raised, unbracketed run exactly where
    it occurred in the original document, rather than every reference in
    the paragraph getting bunched onto the end of it."""
    p = doc.add_paragraph(style=style_name)
    base_size = doc.styles[style_name].font.size
    for kind, value in segments:
        if kind == 'text':
            if value:
                p.add_run(value)
        else:
            marker_run = p.add_run(','.join(str(n) for n in value))
            marker_run.font.superscript = True
            if base_size:
                marker_run.font.size = Pt(round(base_size.pt * NOTE_MARKER_SCALE, 1))
    return p


def _get_or_add_style(doc, name):
    """'Body Text' etc already exist as Word built-ins, so reuse them
    rather than crashing on add_style."""
    for s in doc.styles:
        if s.name == name and s.type == WD_STYLE_TYPE.PARAGRAPH:
            return s
    return doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _segment_word_count(segments):
    """Word count across just the text portions of a segment list -
    note markers shouldn't count toward the intro word budget."""
    return sum(len(value.split()) for kind, value in segments if kind == 'text')


def _build_styles(doc):
    for name, (font, size, leading, bold, italic, caps, justify, space_after_pt) in STYLES.items():
        style = _get_or_add_style(doc, name)
        style.font.name = font
        # Force the eastasian font tag too, otherwise Word/InDesign can
        # silently fall back to a default font for some character ranges.
        style.element.rPr.rFonts.set(qn('w:eastAsia'), font)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.italic = italic
        _set_all_caps(style, caps)
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(leading)
        pf.space_after = Pt(space_after_pt)
        pf.space_before = Pt(0)
        if justify:
            pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return doc


def _extract_note_texts(zf, part_name, tag):
    """Read word/footnotes.xml or word/endnotes.xml and return {id: text},
    skipping the separator/continuation-separator placeholder entries Word
    always includes."""
    if part_name not in zf.namelist():
        return {}
    root = etree.fromstring(zf.read(part_name))
    notes = {}
    for note in root.findall(f'w:{tag}', NSMAP):
        if note.get(f'{{{W_NS}}}type') in ('separator', 'continuationSeparator'):
            continue
        note_id = note.get(f'{{{W_NS}}}id')
        text = ''.join(t.text or '' for t in note.findall('.//w:t', NSMAP)).strip()
        if text:
            notes[note_id] = text
    return notes


def _paragraph_segments(para):
    """Walk a paragraph's runs in document order, splitting it into
    ('text', string) and ('note', (kind, id)) segments. Footnote/endnote
    reference runs in Word don't carry visible text of their own - they
    sit as a distinct run wherever the contributor inserted the
    reference - so walking runs in order (rather than reading
    para.text and a separate list of reference IDs) is what lets a note
    end up attached to the right sentence instead of the end of the
    whole paragraph."""
    segments = []
    buffer = []

    def flush():
        if buffer:
            segments.append(('text', ''.join(buffer)))
            buffer.clear()

    for run in para._p.findall('w:r', NSMAP):
        fn_ref = run.find('w:footnoteReference', NSMAP)
        en_ref = run.find('w:endnoteReference', NSMAP)
        if fn_ref is not None:
            flush()
            segments.append(('note', ('footnote', fn_ref.get(f'{{{W_NS}}}id'))))
        elif en_ref is not None:
            flush()
            segments.append(('note', ('endnote', en_ref.get(f'{{{W_NS}}}id'))))
        else:
            for t in run.findall('w:t', NSMAP):
                buffer.append(t.text or '')
    flush()
    return segments


def _merge_note_segments(segments):
    """Collapse consecutive ('note', ...) segments with nothing but
    other notes between them into a single comma-joined marker, e.g. two
    references back to back render as a single "9,10" superscript run
    rather than two raised digits jammed together unreadably as "910"."""
    merged = []
    pending = []

    def flush_pending():
        if pending:
            merged.append(('note', list(pending)))
            pending.clear()

    for kind, value in segments:
        if kind == 'note':
            pending.append(value)
        else:
            flush_pending()
            merged.append((kind, value))
    flush_pending()
    return merged


def _looks_like_subhead(text):
    """True if a plain, untagged paragraph looks like a section heading
    rather than a sentence of running prose: short, with no terminal
    punctuation, and starting with a capital letter. Deliberately loose
    on capitalisation pattern beyond that - it covers both Title Case
    ("Silicon Economy") and sentence case ("Critique of the family"),
    since real headings show up in both conventions and the punctuation
    + length check is already doing the heavy lifting."""
    words = text.split()
    if not words or len(words) > MAX_SUBHEAD_WORDS:
        return False
    if text[-1] in '.,;:!?':
        return False
    return text[0].isalpha() and text[0].isupper()


def extract_body_paragraphs(filepath):
    """Read the contributor's uploaded .docx and classify each paragraph
    into an IMR style based on [MARKER] tags or Word heading levels.
    Also pulls real Word footnotes/endnotes out of the package directly,
    since python-docx's normal paragraph text silently drops them.
    Returns (paragraphs, notes) where paragraphs is
    [(style_name, segments), ...] - segments is a list of
    ('text', string) / ('note', [numbers]) tuples in document order, so a
    reference renders attached to the exact sentence it belongs to
    rather than bunched onto the end of the paragraph - and notes is
    [(number, text), ...]."""
    src = Document(filepath)
    with zipfile.ZipFile(filepath) as zf:
        footnotes = _extract_note_texts(zf, 'word/footnotes.xml', 'footnote')
        endnotes = _extract_note_texts(zf, 'word/endnotes.xml', 'endnote')

    result = []
    notes = []
    seen = set()

    def register(kind, note_id, source):
        key = (kind, note_id)
        if key in seen or note_id not in source:
            return None
        seen.add(key)
        number = len(notes) + 1
        notes.append((number, source[note_id]))
        return number

    first_body_seen = False
    for para in src.paragraphs:
        full_text = para.text.strip()
        if not full_text:
            continue

        raw_segments = _paragraph_segments(para)
        resolved = []
        for kind, value in raw_segments:
            if kind == 'text':
                resolved.append(('text', value))
            else:
                note_kind, note_id = value
                source = footnotes if note_kind == 'footnote' else endnotes
                n = register(note_kind, note_id, source)
                if n:
                    resolved.append(('note', n))

        style_name = None
        for marker, mapped in MARKER_STYLE.items():
            if full_text.upper().startswith(marker):
                style_name = mapped
                # Markers are plain text the contributor typed, so they
                # land in the first text segment - strip just that
                # prefix rather than touching anything else.
                if resolved and resolved[0][0] == 'text':
                    stripped = resolved[0][1].strip()
                    if stripped.upper().startswith(marker):
                        resolved[0] = ('text', stripped[len(marker):].strip())
                break

        if style_name is None:
            word_style = (para.style.name or '').lower()
            if 'heading' in word_style:
                style_name = 'Subhead'
            elif _looks_like_subhead(full_text):
                style_name = 'Subhead'
            else:
                style_name = 'Body Text First' if not first_body_seen else 'Body Text'
                first_body_seen = True

        result.append((style_name, _merge_note_segments(resolved)))

    return result, notes


def process_docx(filepath, title, author, standfirst, article_type,
                  output_folder, timestamp):
    """Build the IMR-styled output .docx and return its path.

    Layout: page 1 (title, byline, standfirst, a fixed-height blank gap
    reserved for a photo, and a lead chunk of body text up to roughly
    INTRO_WORD_TARGET words) is a single column. From there, everything
    else - subheads, pull quotes, the rest of the body, captions,
    endnotes - flows in two columns starting on page 2, matching the
    print layout of the actual journal.
    """
    out = Document()

    # Strip the default boilerplate styles isn't necessary - we just add ours
    _build_styles(out)
    _set_page_size(out.sections[0])
    _set_margins(out.sections[0])

    out.add_paragraph(title, style='Article Title')
    out.add_paragraph(f'By {author}' if author else '', style='Byline')
    if standfirst:
        out.add_paragraph(standfirst, style='Standfirst')

    _add_fixed_gap(out, IMAGE_GAP_PT)

    body_paragraphs, notes = extract_body_paragraphs(filepath)

    # Build the single-column intro out of plain running body paragraphs
    # until the word budget is hit. A subhead/pull quote/caption always
    # ends the intro immediately, even under budget, since that's the
    # natural break point in print layout.
    intro_paragraphs = []
    remaining_paragraphs = []
    word_count = 0
    in_intro = True
    for style_name, segments in body_paragraphs:
        if in_intro and style_name in INTRO_ELIGIBLE_STYLES:
            intro_paragraphs.append((style_name, segments))
            word_count += _segment_word_count(segments)
            if word_count >= INTRO_WORD_TARGET:
                in_intro = False
        else:
            in_intro = False
            remaining_paragraphs.append((style_name, segments))

    for style_name, segments in intro_paragraphs:
        _add_paragraph_with_notes(out, segments, style_name)

    if remaining_paragraphs or notes:
        two_col_section = out.add_section(WD_SECTION.NEW_PAGE)
        _set_columns(two_col_section, 2)
        _set_page_size(two_col_section)
        _set_margins(two_col_section)

        for style_name, segments in remaining_paragraphs:
            _add_paragraph_with_notes(out, segments, style_name)

        if notes:
            out.add_paragraph('Endnotes', style='Subhead')
            for number, note_text in notes:
                out.add_paragraph(f'{number}. {note_text}', style='Endnote Text')

    safe_title = ''.join(c if c.isalnum() or c in ' -_' else '' for c in title)[:50].strip() or 'untitled'
    filename = f"{timestamp}_{safe_title.replace(' ', '_')}_IMR.docx"
    out_path = os.path.join(output_folder, filename)
    out.save(out_path)
    return out_path
